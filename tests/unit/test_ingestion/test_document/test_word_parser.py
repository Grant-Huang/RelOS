"""
tests/unit/test_ingestion/test_document/test_word_parser.py
-------------------------------------------------------------
Word 解析器单元测试（纯内存，无 Neo4j 依赖）。
"""
from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

from relos.ingestion.document.models import TemplateType
from relos.ingestion.document.word_parser import parse_word


def _make_8d_docx() -> bytes:
    """构造标准 8D 报告 docx 字节。"""
    doc = DocxDocument()
    doc.add_heading("8D质量异常分析报告", level=0)
    doc.add_paragraph("文件编号：QAR-2026-012  日期：2026-03-15")

    doc.add_heading("D1 小组成员", level=1)
    doc.add_paragraph("组长：王工（质量部）  成员：李工（生产部）、张工（设备部）")

    doc.add_heading("D2 问题描述", level=1)
    doc.add_paragraph("产品焊缝宽度超差 0.8mm，不良率 3.2%，批次 B-2026-012。")
    doc.add_paragraph("发现工序：焊接线 L2，发现时间：2026-03-14 14:30。")

    doc.add_heading("D3 临时遏制措施", level=1)
    doc.add_paragraph("立即停止批次 B-2026-012 出货，对在制品全检。")

    doc.add_heading("D4 根本原因", level=1)
    doc.add_paragraph("5-Why 分析：")
    doc.add_paragraph("Why1: 焊缝偏移 → 工件定位不准")
    doc.add_paragraph("Why2: 定位夹具磨损 0.6mm → 累计磨损超限")
    doc.add_paragraph("根本原因：定位夹具磨损超出公差，导致工件偏移，焊缝超出公差带。")

    doc.add_heading("D5 永久纠正措施", level=1)
    doc.add_paragraph("更换所有磨损超限的定位夹具，建立月度点检制度。")

    doc.add_heading("D7 预防措施", level=1)
    doc.add_paragraph("将夹具磨损检查纳入月度 PM 计划，阈值 0.3mm 触发预警。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_shift_docx() -> bytes:
    """构造交接班日志 docx 字节。"""
    doc = DocxDocument()
    doc.add_heading("夜班交接班记录", level=0)

    doc.add_heading("班次信息", level=1)
    doc.add_paragraph("日期：2026-03-20  班次：夜班（22:00-06:00）  交班人：李工")

    doc.add_heading("本班异常事件", level=1)
    doc.add_paragraph("22:15 焊接机 M3 触发过热报警，停机 35 分钟。")
    doc.add_paragraph("处理人：李工，手动降低焊接电流至 180A 后恢复正常。")

    doc.add_heading("遗留问题", level=1)
    doc.add_paragraph("建议白班安排 M3 预防性维护，检查冷却系统。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 8D 报告解析 ──────────────────────────────────────────────────

class TestParse8D:

    def setup_method(self) -> None:
        self.docx_bytes = _make_8d_docx()

    def test_template_type(self) -> None:
        doc = parse_word(self.docx_bytes, "8d_report.docx")
        assert doc.template_type == TemplateType.QUALITY_8D

    def test_has_sections(self) -> None:
        doc = parse_word(self.docx_bytes, "8d_report.docx")
        assert len(doc.sections) > 0

    def test_d4_root_cause_extracted(self) -> None:
        doc = parse_word(self.docx_bytes, "8d_report.docx")
        # D4 章节应包含根因信息
        all_content = " ".join(doc.sections.values())
        assert "夹具" in all_content or "根本原因" in all_content

    def test_d2_problem_extracted(self) -> None:
        doc = parse_word(self.docx_bytes, "8d_report.docx")
        all_content = " ".join(doc.sections.values())
        assert "焊缝" in all_content or "超差" in all_content

    def test_source_filename(self) -> None:
        doc = parse_word(self.docx_bytes, "my_8d.docx")
        assert doc.source_filename == "my_8d.docx"

    def test_no_rows(self) -> None:
        doc = parse_word(self.docx_bytes, "8d.docx")
        assert doc.rows == []


# ─── 交接班记录解析 ───────────────────────────────────────────────

class TestParseShiftHandover:

    def setup_method(self) -> None:
        self.docx_bytes = _make_shift_docx()

    def test_template_type(self) -> None:
        doc = parse_word(self.docx_bytes, "shift.docx")
        assert doc.template_type == TemplateType.SHIFT_HANDOVER

    def test_abnormal_event_extracted(self) -> None:
        doc = parse_word(self.docx_bytes, "shift.docx")
        all_content = " ".join(doc.sections.values())
        assert "M3" in all_content or "过热" in all_content

    def test_pending_issue_extracted(self) -> None:
        doc = parse_word(self.docx_bytes, "shift.docx")
        all_content = " ".join(doc.sections.values())
        assert "维护" in all_content or "冷却" in all_content


# ─── 无效文件 ─────────────────────────────────────────────────────

class TestInvalidFile:

    def test_invalid_file_raises(self) -> None:
        with pytest.raises(ValueError, match="无法读取"):
            parse_word(b"not a docx file", "bad.docx")
